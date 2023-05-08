###简单的提问方式
# 使用 OpenAI API 向其提问
import openai
import os
OPENAI_API_KEY = "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"
openai.api_key = OPENAI_API_KEY

# 使用 OpenAI API 向其提问
def ask_openai(question):
   prompt = (f"Answer the following question: \n"
             f"{question}\n"
             f"Answer:")
   messages = [{"role": "user", "content": prompt}]

   response = openai.ChatCompletion.create(
   messages=messages, max_tokens=2048, n=1,stop=None,temperature=0.5,model="gpt-3.5-turbo"
)

   answer = response.choices[0].message["content"]
   return answer

question = "我向你提问的上一个问题是什么？"
answer = ask_openai(question)
print(f"Question: {question}\nAnswer: {answer}")
-----------------------------------------------------------------------------

###让gpt概括内容
import openai
import os
OPENAI_API_KEY = "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"
openai.api_key = OPENAI_API_KEY

def get_completion(prompt, model="gpt-3.5-turbo"): 
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        temperature=0, # 值越低则输出文本随机性越低
    )
    return response.choices[0].message["content"]
   
   
   -----------------------------------------------------------------------
import docx
import openai

# 设置 OpenAI API 密钥
openai.api_key = "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"

# 读取 Word 文档内容
doc = docx.Document("/Users/long/Desktop/五一新闻.docx")
text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

# print(text)

def get_completion(prompt, model="gpt-3.5-turbo"):
    '''
    prompt: 对应的提示
    model: 调用的模型，默认为 gpt-3.5-turbo(ChatGPT)，有内测资格的用户可以选择 gpt-4
    '''
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        n=1,
        stop=None,
        frequency_penalty=0,
        presence_penalty=0,
        max_tokens=2048,
        temperature=0.7, # 模型输出的温度系数，控制输出的随机程度
    )
    # 调用 OpenAI 的 ChatCompletion 接口
    return response.choices[0].message["content"]

while True:
# 提示输入问题
  question = input("请输入你的问题：")
#   prompt = f"Q: {question}\nA:"
  prompt = f"请围绕下面的文本内容，根据问题，做出回答。\
    文本内容是：{text}\
    问题是：{question}\
        输出格式为：Q：<>\nA:<>"""
  
  response = get_completion(prompt)
  print(response)

--------------------------------------------------------------
import docx
import openai

# 设置 OpenAI API 密钥
openai.api_key = "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"

# 读取 Word 文档内容
doc = docx.Document("/Users/long/Desktop/五一新闻.docx")
text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

# print(text)

def get_completion(prompt, model="gpt-3.5-turbo"):
    '''
    prompt: 对应的提示
    model: 调用的模型，默认为 gpt-3.5-turbo(ChatGPT)，有内测资格的用户可以选择 gpt-4
    '''
    messages = [{"role": "user", "content": prompt}]
    response = openai.ChatCompletion.create(
        model=model,
        messages=messages,
        n=1,
        stop=None,
        frequency_penalty=0,
        presence_penalty=0,
        max_tokens=2048,
        temperature=0.5, # 模型输出的温度系数，控制输出的随机程度
    )
    # 调用 OpenAI 的 ChatCompletion 接口
    return response.choices[0].message["content"]

while True:
# 提示输入问题
  question = input("请输入你的问题：")
#   prompt = f"Q: {question}\nA:"
  prompt = f"请结合下面的文本内容，根据你自己的经验知识，总结文本内容表现出的管理现状情况，并给出初步的管理建议。（管理现状和初步建议要换行区分，分别表达；管理现状和初步建议中用序号分开进行要点表达，但要点要尽量抽象概括，不要只是复述数据）不要再把数据复述一遍，直接给出概括性的结论。\
    文本内容是：{text}\
    问题是：{question}\
    """
  response = get_completion(prompt)
  report =  f"一、根据目前提供的指标对现状进行分析，给出初步建议：\n   {response}\n\n二、行业经验表明，从「产品准确度」、「质量」、「交付速度」、「研发团队个人效能」几个维度对研发效能进行考察是最佳实践。您提供的指标尚不全面，建议在四个维度上补充更多指标，以帮助我站在更全面的视角为您提供更准确的现状分析和管理建议。下面是关于效能四维度的解释，以及各个维度的相关指标，并附带了详细的指标相关业务数据的搜集策略和统计方法。\n  1. 产品准确度\n产品准确度是指产品能否满足用户需求的程度。为了提高产品准确度，企业可以采用用户体验设计（UX）方法，进行用户研究、需求分析和交互设计等，以确保产品能够满足用户的期望。同时，对产品进行严格的质量控制和测试，以确保产品能够稳定运行并符合功能要求。\n相关指标模型统计方法和数据搜集策略：www.baidu.com\n\n2. 质量\n产品质量是指产品能够满足用户期望的程度。为了提高产品质量，企业可以采用质量管理（QM）方法，如Six Sigma和ISO 9001等，对产品的整个生命周期进行管理和控制，以确保产品的稳定性和可靠性。\n相关指标模型统计方法和数据搜集策略：www.baidu.com\n\n3. 交付速度\n交付速度是指企业向客户交付产品和服务的速度。为了提高交付速度，企业可以采用敏捷开发（Agile Development）和持续集成（Continuous Integration）等方法，以便更快地推出新产品和服务，并更快地满足客户需求。\n相关指标模型统计方法和数据搜集策略：www.baidu.com\n\n4. 研发团队个人效能\n研发团队个人效能是指团队成员在研发过程中的效率和能力。为了提高研发团队个人效能，企业可以采用各种培训和发展计划，以提高团队成员的技能和知识水平。此外，企业还可以通过激励计划和绩效评估，激励团队成员在研发过程中做出更好的表现。\n相关指标模型统计方法和数据搜集策略：www.baidu.com\n\n\n三、当您的团队在不断发展壮大时，如何高效地管理团队效能是至关重要的问题。ONES是一款高效的团队效能管理工具，它可以帮助您实现全方位的团队效能管理，包括任务管理、进度跟踪、协作沟通、绩效评估等多个方面。\n\n欲了解ONES更多信息，请移步官网：www.baidu.com"
  

  print(report)

